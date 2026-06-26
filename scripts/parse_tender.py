#!/usr/bin/env python3
"""
解析招标技术规格书（docx），输出结构化层级树 JSON。

功能：
1. 提取段落文本和格式（加粗、字号、颜色）
2. 识别章节层级（基于大纲级别、字号、编号模式）
3. 检测重点标记（加粗、标星、标色、"重点/关键/核心"关键词）
4. 输出 JSON 树结构
"""

import json
import re
import sys
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("错误：请先安装 python-docx：pip install python-docx", file=sys.stderr)
    sys.exit(1)


def get_outline_level(paragraph) -> Optional[int]:
    """获取段落的大纲级别 (1-9)，None 表示正文"""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is not None:
            return int(outline.get(qn("w:val"))) + 1
    return None


def get_run_format_info(run) -> dict:
    """提取单个 run 的格式信息"""
    info: dict[str, Any] = {}
    info["bold"] = run.bold  # True/False/None
    info["italic"] = run.italic

    if run.font.size:
        info["font_size_pt"] = run.font.size.pt
    if run.font.color and run.font.color.rgb:
        info["color"] = str(run.font.color.rgb)

    # 检测高亮
    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        highlight = rPr.find(qn("w:highlight"))
        if highlight is not None:
            info["highlight"] = highlight.get(qn("w:val"))

    return info


def detect_emphasis(text: str, run_info: dict, all_run_infos: list[dict]) -> dict:
    """检测重点标记程度，返回 emphasis 对象"""
    score = 0
    signals = []

    # 加粗
    if run_info.get("bold"):
        score += 30
        signals.append("bold")

    # 标色
    color = run_info.get("color", "")
    if color and color.upper() not in ("000000", "333333", "FFFFFF", "AUTO"):
        score += 25
        signals.append(f"color#{color}")

    # 高亮
    if run_info.get("highlight") and run_info.get("highlight") not in ("none", "auto"):
        score += 25
        signals.append("highlighted")

    # 标星/特殊符号
    star_pattern = r"[★☆※●◦▪▸►❖◆◇▲△▶▷✓✔☑✅]"
    if re.search(star_pattern, text):
        score += 20
        signals.append("symbol")

    # 关键词
    emphatic_keywords = [
        "重点", "关键", "核心", "必须", "强制性", "否决项",
        "废标项", "实质性要求", "重要"
    ]
    found_kw = [kw for kw in emphatic_keywords if kw in text]
    if found_kw:
        score += 20
        signals.append(f"keyword:{','.join(found_kw)}")

    return {
        "score": min(score, 100),
        "is_emphasized": score >= 30,  # 阈值为 30 分即视为重点
        "signals": signals,
    }


def detect_numbering(text: str) -> Optional[dict]:
    """检测标题/条款编号模式"""
    patterns = [
        # 中文序号：第X章、第一节
        (r"^第[一二三四五六七八九十百千\d]+[章节条款].*", "chapter_section"),
        # 数字层级：3.1.2、4.2.1.3
        (r"^(\d+[.、]){1,4}\d*\s*.*", "numeric"),
        # 中文数字：一、二、三、
        (r"^[一二三四五六七八九十]+[、，,]\s*.*", "cn_num"),
        # 带括号：1)、（一）、(1)
        (r"^[\(（]\s*[\d一二三四五六七八九十]+\s*[\)）].*", "parenthesized"),
        # 编号+空格：1 概述
        (r"^\d+\s+\S", "number_space"),
    ]
    for pat, style in patterns:
        if re.match(pat, text.strip()):
            return {"style": style, "pattern": pat}
    return None


def extract_headings(doc: Document) -> list[dict]:
    """从 Word 内置标题样式中提取标题，建立标题索引"""
    headings = []
    # 收集所有带标题样式的段落
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level_str = para.style.name.replace("Heading", "").strip()
            try:
                level = int(level_str)
            except ValueError:
                level = 1
            headings.append({
                "text": para.text.strip(),
                "level": level,
                "para_index": i,
                "style_name": para.style.name,
            })
    return headings


def parse_tender(filepath: str) -> dict:
    """主解析函数"""
    doc = Document(filepath)
    result: dict[str, Any] = {
        "file": str(Path(filepath).absolute()),
        "total_paragraphs": len(doc.paragraphs),
        "sections": [],      # 顶级章节
        "flat_headings": [], # 平铺标题列表（含层级）
        "emphasized_items": [],  # 重点标记条目
    }

    # 先提取 Word 内置标题作为骨架
    word_headings = extract_headings(doc)

    # 遍历段落，构建结构
    current_section = None  # 顶级章节
    section_stack: list[dict] = []  # 章节栈

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 获取大纲级别
        outline_lvl = get_outline_level(para)

        # 获取段落中所有 run 的格式
        runs_info = []
        full_emphasis = {"score": 0, "is_emphasized": False, "signals": []}
        for run in para.runs:
            ri = get_run_format_info(run)
            runs_info.append(ri)
            emp = detect_emphasis(run.text, ri, runs_info)
            if emp["score"] > full_emphasis["score"]:
                full_emphasis = emp

        # 二次检测：对整段文本做关键词检测
        text_emp = detect_emphasis(text, {}, [])
        if text_emp["score"] > full_emphasis["score"]:
            full_emphasis = text_emp

        # 检测编号
        numbering = detect_numbering(text)

        # 判断是否为标题
        is_heading = bool(
            outline_lvl is not None
            or numbering
            or (runs_info and runs_info[0].get("bold") and len(text) < 80)
            or any(h["para_index"] == i for h in word_headings)
        )

        # 判断层级
        if outline_lvl is not None:
            level = outline_lvl
        elif numbering:
            # 根据编号模式推断层级
            if numbering["style"] in ("chapter_section", "cn_num"):
                level = 1
            elif numbering["style"] == "numeric":
                # 点数越多层级越深：3.1.2 = level 3
                depth = len(re.findall(r"\d+", text.split()[0] if text.split() else ""))
                level = min(depth, 3) if depth > 0 else 2
            else:
                level = 2
        elif is_heading:
            level = 2
        else:
            level = 0  # 正文

        # 构建条目
        entry = {
            "index": i,
            "text": text,
            "level": level,
            "is_heading": is_heading,
            "outline_level": outline_lvl,
            "numbering": numbering,
            "emphasis": full_emphasis,
        }

        # 添加到扁平列表
        result["flat_headings"].append(entry)

        # 重点标记收集
        if full_emphasis["is_emphasized"]:
            result["emphasized_items"].append(entry)

        # 构建层级树
        if not section_stack:
            section_stack.append({
                "title": text,
                "level": 1,
                "children": [],
                "emphasis": full_emphasis,
                "para_index": i,
            })
            result["sections"].append(section_stack[0])
        else:
            # 简化处理：根据 level 调整栈
            while len(section_stack) > 1 and section_stack[-1]["level"] >= level:
                section_stack.pop()

            if level > section_stack[-1]["level"]:
                child = {
                    "title": text,
                    "level": level,
                    "children": [],
                    "emphasis": full_emphasis,
                    "para_index": i,
                }
                section_stack[-1]["children"].append(child)
                section_stack.append(child)
            else:
                sibling = {
                    "title": text,
                    "level": level,
                    "children": [],
                    "emphasis": full_emphasis,
                    "para_index": i,
                }
                if len(section_stack) > 1:
                    section_stack[-2]["children"].append(sibling)
                    section_stack[-1] = sibling
                else:
                    result["sections"].append(sibling)
                    section_stack[0] = sibling

    # 统计
    result["heading_count"] = len(result["flat_headings"])
    result["emphasized_count"] = len(result["emphasized_items"])

    return result


def format_sections(sections: list[dict], indent: int = 0) -> str:
    """格式化输出章节树"""
    lines = []
    prefix = "  " * indent
    for sec in sections:
        emp = sec.get("emphasis", {})
        markers = []
        if emp.get("is_emphasized"):
            signals = emp.get("signals", [])
            if "bold" in signals:
                markers.append("🔴加粗")
            if any("color" in s for s in signals):
                markers.append("🎨标色")
            if "keyword" in ",".join(signals):
                markers.append("⭐重点词")
            if "symbol" in signals:
                markers.append("★符号")
        marker_str = f" [{' '.join(markers)}]" if markers else ""
        lines.append(f"{prefix}├─ {sec['title'][:80]}{marker_str}")
        for child in sec.get("children", []):
            lines.append(format_sections([child], indent + 1))
    return "\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print("用法: python parse_tender.py <招标文件.docx>")
        print("输出: JSON 结构化解析结果（stdout）")
        sys.exit(1)

    filepath = sys.argv[1]
    if not Path(filepath).exists():
        print(f"错误：文件不存在 — {filepath}", file=sys.stderr)
        sys.exit(1)

    result = parse_tender(filepath)

    # 输出完整 JSON
    if "--json" in sys.argv:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        # 友好输出模式：先输出章节树，再输出 JSON
        print("=" * 60)
        print(f"📄 文件：{result['file']}")
        print(f"📊 共 {result['total_paragraphs']} 段，识别 {result['heading_count']} 个标题")
        print(f"⭐ 重点标记条目：{result['emphasized_count']} 个")
        print("=" * 60)
        print("\n📑 章节结构：\n")
        print(format_sections(result["sections"]))
        print("\n" + "=" * 60)
        print("\n⭐ 重点标记条目详情：\n")
        for item in result["emphasized_items"]:
            sigs = item["emphasis"].get("signals", [])
            print(f"  [{item['index']}] {item['text'][:100]}")
            print(f"       标记：{', '.join(sigs)} (score={item['emphasis']['score']})")
            print()

        # 仍然输出 JSON 到 stdout 供后续管道使用
        json_str = json.dumps(result, ensure_ascii=False, indent=2)
        print("\n--- JSON OUTPUT ---")
        print(json_str)


if __name__ == "__main__":
    main()
