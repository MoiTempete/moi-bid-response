#!/usr/bin/env python3
"""
生成投标技术响应 Word 文档 — 结构化技术文档 + 国标多级自动编号。

输入：JSON（从文件或 stdin 读取，格式见 SKILL.md Step 5）
      - 推荐方式：python3 generate_docx.py output.docx data.json
      - 兼容方式：echo '<JSON>' | python3 generate_docx.py output.docx
输出：格式化的 docx 文件

特性和格式（匹配往期2真实投标文件）：
- A4 页面，左/右边距 3.2cm，上/下边距 2.5cm
- 宋体正文 12pt，1.5 倍行距
- 黑体标题（H2~H5），1.5 倍行距
- 国标多级自动编号：一、/ 1. / 1.1. / 1.1.1. / ① / (1)
- 嵌入式表格（宋体 11pt，黑体表头加粗灰底）
- 无页眉页脚（简洁风格）
"""

import json
import sys
from datetime import datetime
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
except ImportError:
    print("错误：请先安装 python-docx：pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 格式常量（匹配往期2真实格式）
# ============================================================

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_SIZE_BODY = Pt(12)          # 正文 12pt
FONT_SIZE_TABLE = Pt(11)         # 表格 11pt
FONT_SIZE_H1 = Pt(16)            # 文档标题 16pt
FONT_SIZE_H2 = Pt(15)            # 一级目录 15pt
FONT_SIZE_H3 = Pt(14)            # 二级目录 14pt
FONT_SIZE_H4 = Pt(14)            # 三级目录及以下 14pt
LINE_SPACING = 1.5               # 全局行距
HEADING_SPACE_BEFORE = Pt(0)     # 标题段前距（不加额外间距）
HEADING_SPACE_AFTER = Pt(0)      # 标题段后距（不加额外间距）
BODY_SPACE_BEFORE = Pt(0)        # 正文段前距（不加额外间距）
BODY_SPACE_AFTER = Pt(0)         # 正文段后距（不加额外间距）
BODY_INDENT = Cm(0.74)           # 正文首行缩进（约2个中文字符）

TABLE_HEADER_BG = "D9D9D9"       # 表头灰底

# 页边距
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3.2)
MARGIN_RIGHT = Cm(3.2)


# ============================================================
# 样式配置
# ============================================================

def setup_styles(doc: Document):
    """配置文档默认样式"""
    # Normal（正文）
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = BODY_SPACE_BEFORE
    pf.space_after = BODY_SPACE_AFTER

    # Heading 1（文档标题）
    style = doc.styles["Heading 1"]
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_H1
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)

    # Heading 2-5（各级目录）
    for i, size in [(2, FONT_SIZE_H2), (3, FONT_SIZE_H3), (4, FONT_SIZE_H4), (5, FONT_SIZE_H4)]:
        try:
            style = doc.styles[f"Heading {i}"]
        except KeyError:
            continue
        style.font.name = FONT_HEADING
        style.font.size = size
        style.font.bold = False
        style.font.color.rgb = None  # 黑色
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = HEADING_SPACE_BEFORE
        pf.space_after = HEADING_SPACE_AFTER
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.keep_with_next = True

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


# ============================================================
# 国标多级自动编号（核心）
# ============================================================

def setup_multilevel_numbering(doc: Document):
    """
    设置 Word 多级列表编号方案：

    H1: 无编号（文档标题）
    H2: 一、二、三、...          (chineseCountingThousand + 、)
    H3: 1. 2. 3. ...            (decimal + .)
    H4: 1.1. 1.2. ...           (decimal + .)
    H5: 1.1.1. 1.1.2. ...       (decimal + .)
    H6: ①②③...（尽力支持）       (decimal + 特殊 lvlText)
    H7: (1)(2)(3)...（尽力支持）  (decimal + 括号)
    """
    # 创建抽象编号定义
    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), "0")
    abstractNum.set(qn("w:restartNumberingAfterBreak"), "0")

    # 多级类型
    multiLevelType = OxmlElement("w:multiLevelType")
    multiLevelType.set(qn("w:val"), "hybridMultilevel")
    abstractNum.append(multiLevelType)

    # 定义各层级的编号格式
    # OOXML 中 lvlText 的 %N 规则：
    #   %1 = ilvl 0 的编号, %2 = ilvl 1 的编号, %3 = ilvl 2 的编号, ...
    #   当前层级可引用 ≤ 当前 ilvl+1 编号
    #
    # 国标方案：
    #   ilvl 0 (H2): 一、二、三、         ← %1、 (chineseCounting)
    #   ilvl 1 (H3): 1. 2. 3.           ← %2.  (独立编号，在 H2 下重启)
    #   ilvl 2 (H4): 1.1. 1.2.          ← %2.%3. (显示 H3.H4)
    #   ilvl 3 (H5): 1.1.1. 1.1.2.      ← %2.%3.%4.
    #   ilvl 4 (H6): ①②③               ← ① (特殊，编号文本固定)
    #   ilvl 5 (H7): (1)(2)(3)          ← (%7) (特殊)
    #
    # 注意：格式 ① 和 (N) 在 Word 中无法完美自动递增，脚本尽力支持。
    #
    levels_config = [
        (0, "chineseCountingThousand", "%1、", "1",
         {"left": Cm(0), "hanging": Cm(1.2)}),
        (1, "decimal", "%2.", "1",
         {"left": Cm(0.8), "hanging": Cm(1.2)}),
        (2, "decimal", "%2.%3.", "1",
         {"left": Cm(1.6), "hanging": Cm(1.6)}),
        (3, "decimal", "%2.%3.%4.", "1",
         {"left": Cm(2.4), "hanging": Cm(2.0)}),
        (4, "decimal", "①", "1",
         {"left": Cm(3.2), "hanging": Cm(1.5)}),
        (5, "decimal", "(%7)", "1",
         {"left": Cm(4.0), "hanging": Cm(1.5)}),
    ]

    for ilvl, numFmt, lvlText, start, indent in levels_config:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        lvl.set(qn("w:tplc"), "FFFFFFFF")

        # 起始值
        start_elem = OxmlElement("w:start")
        start_elem.set(qn("w:val"), start)
        lvl.append(start_elem)

        # 编号格式
        numFmt_elem = OxmlElement("w:numFmt")
        numFmt_elem.set(qn("w:val"), numFmt)
        lvl.append(numFmt_elem)

        # 编号文本
        lvlText_elem = OxmlElement("w:lvlText")
        lvlText_elem.set(qn("w:val"), lvlText)
        lvl.append(lvlText_elem)

        # 对齐和缩进（1cm ≈ 567 twips）
        pPr = OxmlElement("w:pPr")
        ind_elem = OxmlElement("w:ind")
        left_twips = int(indent["left"] / Cm(1) * 567)
        hanging_twips = int(indent["hanging"] / Cm(1) * 567)
        ind_elem.set(qn("w:left"), str(left_twips))
        ind_elem.set(qn("w:hanging"), str(hanging_twips))
        pPr.append(ind_elem)
        lvl.append(pPr)

        abstractNum.append(lvl)

    # 创建编号实例，关联到抽象编号
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")
    abstractNumId_ref = OxmlElement("w:abstractNumId")
    abstractNumId_ref.set(qn("w:val"), "0")
    num.append(abstractNumId_ref)

    # 将编号定义注入文档的 numbering 部分
    numbering_part = doc.part.numbering_part
    numbering = numbering_part._element

    # 清除旧的 abstractNum 和 num，避免与模板默认编号冲突
    for old_an in numbering.findall(qn("w:abstractNum")):
        numbering.remove(old_an)
    for old_num in numbering.findall(qn("w:num")):
        numbering.remove(old_num)

    numbering.append(abstractNum)
    numbering.append(num)

    # 将 heading 样式关联到编号层级
    # H2(level 2) → ilvl 0, H3(level 3) → ilvl 1, H4 → ilvl 2, H5 → ilvl 3
    heading_to_ilvl = {2: 0, 3: 1, 4: 2, 5: 3, 6: 4, 7: 5}


def add_heading_with_numbering(doc: Document, text: str, level: int):
    """
    添加带自动编号的标题段落。
    通过 numPr 关联到多级列表。

    level 2-7 对应 H2-H7（即 Word 的 Heading 2-7 样式）。
    H1 不编号，直接使用 Heading 1 样式。
    """
    para = doc.add_paragraph(style=f"Heading {level}")

    if level >= 2:
        # 关联到多级列表编号
        heading_to_ilvl = {2: 0, 3: 1, 4: 2, 5: 3, 6: 4, 7: 5}
        ilvl = heading_to_ilvl.get(level, level - 2)

        pPr = para._element.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "1")
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), str(ilvl))
        numPr.append(ilvl_elem)
        pPr.append(numPr)

    # 添加标题文字。注意：不能加 run 覆盖编号，编号由 Word 自动渲染
    # 需要清空默认的编号 run，然后添加标题文字 run
    # 先用空 run 占位（编号位置），再添加文本 run
    run = para.add_run(text)
    # run 的字体继承段落样式（已在 setup_styles 中定义）

    return para


def add_body_paragraph(doc: Document, text: str):
    """添加正文段落"""
    para = doc.add_paragraph(style="Normal")
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = FONT_SIZE_BODY
    return para


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """添加格式化表格（表头灰底黑体加粗，正文宋体）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = FONT_HEADING
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        run.font.size = FONT_SIZE_TABLE
        # 灰底
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_HEADER_BG)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            run.font.size = FONT_SIZE_TABLE

    # 添加段后间距
    doc.add_paragraph()

    return table


# ============================================================
# 递归渲染内容
# ============================================================

def render_content(doc: Document, content_items: list[dict]):
    """递归渲染 content 数组：段落 / 子章节 / 表格"""
    for item in content_items:
        item_type = item.get("type", "")

        if item_type == "p":
            add_body_paragraph(doc, item["text"])

        elif item_type == "table":
            add_table(doc, item.get("headers", []), item.get("rows", []))

        elif "heading" in item and "level" in item:
            # 嵌套子章节
            level = item["level"]
            heading = item["heading"]
            children = item.get("content", [])

            add_heading_with_numbering(doc, heading, level)
            render_content(doc, children)

        elif item_type == "" and "heading" not in item:
            # 未知类型，当作段落处理
            if "text" in item:
                add_body_paragraph(doc, item["text"])


# ============================================================
# 主函数
# ============================================================

def generate_response_docx(data: dict, output_path: str):
    """主生成函数"""
    doc = Document()

    # 1. 样式初始化
    setup_styles(doc)

    # 2. 设置多级自动编号
    setup_multilevel_numbering(doc)

    # 3. 文档标题 (H1，不编号)
    title = data.get("title", "技术响应方案")
    doc.add_heading(title, level=1)

    # 4. 渲染各章节
    sections = data.get("sections", [])
    for section in sections:
        heading = section.get("heading", "")
        level = section.get("level", 2)
        content = section.get("content", [])

        if heading:
            add_heading_with_numbering(doc, heading, level)
        render_content(doc, content)

    # 5. 保存
    doc.save(output_path)
    print(f"✅ 响应文件已生成：{output_path}")
    print(f"   格式：A4 / 宋体12pt正文 / 黑体标题 / 国标多级自动编号 / 1.5倍行距")


def main():
    if len(sys.argv) < 2:
        print("用法:")
        print("  方式1（从文件读取）: python3 generate_docx.py <输出文件名.docx> <JSON文件.json>")
        print("  方式2（从stdin读取）: echo '<JSON>' | python3 generate_docx.py <输出文件名.docx>")
        print()
        print("生成 Word 文档，格式见 SKILL.md Step 5。")
        sys.exit(1)

    output_path = sys.argv[1]
    json_file = sys.argv[2] if len(sys.argv) >= 3 else None

    # 读取 JSON（优先从文件，fallback 到 stdin）
    if json_file:
        # 方式1：从 JSON 文件读取（推荐，避免管道编码问题）
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except FileNotFoundError:
            print(f"错误：JSON 文件不存在 — {json_file}", file=sys.stderr)
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"错误：JSON 解析失败 — {e}", file=sys.stderr)
            sys.exit(1)
    else:
        # 方式2：从 stdin 读取（兼容旧用法）
        raw = sys.stdin.read()
        if not raw:
            print("错误：stdin 无数据，也未指定 JSON 文件路径", file=sys.stderr)
            sys.exit(1)
        try:
            data = json.loads(raw)
        except json.JSONDecodeError as e:
            print(f"错误：JSON 解析失败 — {e}", file=sys.stderr)
            sys.exit(1)

    generate_response_docx(data, output_path)


if __name__ == "__main__":
    main()
